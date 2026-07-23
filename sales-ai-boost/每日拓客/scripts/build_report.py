#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
build_report.py —— 把调研报告 markdown 转 docx，并向汇总表 xlsx 追加本批行。

用法：
  python build_report.py --md reports_2026-07-23.md \
                         --docx 公司调研报告汇总_2026-07-23.docx \
                         --xlsx 拓客汇总表.xlsx \
                         --rows rows.csv

rows.csv 每行（不含表头）：
  序号,拓客日期,公司/集团名,所处城市,具体区域,行业,细分行业,具体情况,主要优势,公司人数

依赖：pip install python-docx openpyxl
"""
import argparse
import csv
import os

from docx import Document
from docx.shared import Pt, RGBColor
import openpyxl
import re


def md_to_docx(md_path, docx_path):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Microsoft YaHei'
    st.font.size = Pt(10.5)

    def add_runs(p, text):
        for part in re.split(r'(\*\*.*?\*\*)', text):
            if part.startswith('**') and part.endswith('**') and len(part) > 4:
                r = p.add_run(part[2:-2])
                r.bold = True
            else:
                p.add_run(part)

    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    first = True
    for ln in lines:
        s = ln.strip()
        if s == '' or s == '---':
            continue
        if s.startswith('# '):
            if first:
                doc.add_heading(s[2:], 0)
                first = False
            else:
                doc.add_heading(s[2:], 1)
        elif s.startswith('### '):
            doc.add_heading(s[4:], 2)
        elif s.startswith('> '):
            p = doc.add_paragraph()
            r = p.add_run(s[2:])
            r.italic = True
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        elif s.startswith('**') and s.endswith('**'):
            p = doc.add_paragraph()
            r = p.add_run(s[2:-2])
            r.bold = True
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
    doc.save(docx_path)
    return sum(1 for l in lines if l.strip().startswith('### '))


def append_xlsx(xlsx_path, rows_csv, tier_map=None):
    rows = []
    with open(rows_csv, encoding='utf-8') as f:
        for r in csv.reader(f):
            if r:
                rows.append([c.strip() for c in r])
    if not rows:
        return 0
    if os.path.exists(xlsx_path):
        wb = openpyxl.load_workbook(xlsx_path)
        ws = wb.active
    else:
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(['序号', '拓客日期', '公司/集团名', '所处城市', '具体区域',
                   '行业', '细分行业', '具体情况', '主要优势', '公司人数',
                   '价值层级'])
    has_tier = ws.max_column >= 11 and ws.cell(1, 11).value == '价值层级'
    if not has_tier:
        # 老表无价值层级列，补一列表头
        ws.cell(1, ws.max_column + 1, '价值层级')
    for r in rows:
        if tier_map is not None:
            name = r[2] if len(r) > 2 else ''
            r = r + [tier_map.get(name, '')]
        ws.append(r)
    wb.save(xlsx_path)
    return len(rows)


def load_tier(path):
    """读取 tier.csv（两列 公司名,层级，无表头）→ {公司名: 层级}。"""
    m = {}
    with open(path, encoding='utf-8') as f:
        for r in csv.reader(f):
            if len(r) >= 2:
                m[r[0].strip()] = r[1].strip()
    return m


def main():
    ap = argparse.ArgumentParser(description='报告 markdown→docx + 汇总表追加')
    ap.add_argument('--md', required=True, help='调研报告 markdown 源文件')
    ap.add_argument('--docx', default='公司调研报告汇总.docx', help='输出 docx 路径')
    ap.add_argument('--xlsx', default='拓客汇总表.xlsx', help='汇总表 xlsx 路径')
    ap.add_argument('--rows', required=True, help='本批行 CSV（10 列，无表头）')
    ap.add_argument('--tier', default=None,
                    help='价值层级 CSV（两列 公司名,层级，无表头）；按公司名匹配写入末列')
    args = ap.parse_args()

    tier_map = load_tier(args.tier) if args.tier else None
    n = md_to_docx(args.md, args.docx)
    m = append_xlsx(args.xlsx, args.rows, tier_map)
    print('docx 公司数=%d, xlsx 追加行数=%d' % (n, m))


if __name__ == '__main__':
    main()
